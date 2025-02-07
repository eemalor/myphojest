from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

app = Flask(__name__)

def set_margins(doc):
    """Устанавливает поля документа."""
    section = doc.sections[0]
    section.top_margin = Inches(0.79)  # 2 см
    section.bottom_margin = Inches(0.79)  # 2 см
    section.left_margin = Inches(1.18)  # 3 см
    section.right_margin = Inches(0.59)  # 1.5 см

def set_font_and_spacing(paragraph):
    """Устанавливает шрифт, размер и межстрочный интервал."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5  # Полуторный интервал
    paragraph_format.alignment = 4  # Выравнивание по ширине
    paragraph_format.first_line_indent = Inches(0.49)  # Абзацный отступ 1.25 см
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

def add_title_page(doc, topic):
    """Добавляет титульный лист."""
    doc.add_heading('Итоговый проект', level=0)
    doc.add_heading(topic, level=1)
    doc.add_paragraph('Выполнил: Фамилия Имя, 9 класс')
    doc.add_paragraph('Руководитель: Фамилия Имя, учитель')
    doc.add_paragraph('Год: 2023')
    doc.add_page_break()

def add_table_of_contents(doc):
    """Добавляет оглавление."""
    doc.add_heading('Содержание', level=1)
    # Добавляем автоматическое оглавление
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    doc.add_page_break()

def generate_project(topic, plan):
    """Генерирует проект."""
    doc = Document()
    set_margins(doc)

    # Титульный лист
    add_title_page(doc, topic)

    # Оглавление
    add_table_of_contents(doc)

    # Введение
    doc.add_heading('Введение', level=1)
    intro = doc.add_paragraph('Введение должно содержать обоснование выбора темы, актуальность, цель и задачи проекта.')
    set_font_and_spacing(intro)

    # Основная часть
    doc.add_heading('Основная часть', level=1)
    main_part = doc.add_paragraph('Основная часть включает теоретическую и практическую части. Здесь вы описываете свои исследования и результаты.')
    set_font_and_spacing(main_part)

    # Заключение
    doc.add_heading('Заключение', level=1)
    conclusion = doc.add_paragraph('Заключение содержит выводы по проекту, рекомендации и возможные перспективы дальнейшего исследования.')
    set_font_and_spacing(conclusion)

    # Список литературы
    doc.add_heading('Список литературы', level=1)
    literature = doc.add_paragraph('1. Автор, Название книги, Год издания.\n2. Автор, Название статьи, Год публикации.')
    set_font_and_spacing(literature)

    # Приложения (если есть)
    doc.add_heading('Приложения', level=1)
    appendix = doc.add_paragraph('Приложения могут включать дополнительные материалы: таблицы, графики, фотографии и т.д.')
    set_font_and_spacing(appendix)

    return doc

@app.route('/generate', methods=['POST'])
def generate():
    data = request.json
    topic = data.get('topic')
    plan = data.get('plan')

    # Генерация документа
    doc = generate_project(topic, plan)

    # Сохранение в BytesIO для отправки
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(file_stream, as_attachment=True, download_name='project.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)